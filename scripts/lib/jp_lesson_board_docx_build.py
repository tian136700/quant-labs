#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""日语新课板书 Word：词卡行切图 + 图下对齐读音表（OJAD）。

依赖：Pillow、python-docx（见 setup 脚本里的 venv）。
算法与 src/lib/jp-vocab-ref-card-grid-crop.ts / check_vocab_ref_card_grid_crop.py 一致。
"""

from __future__ import annotations

import hashlib
import io
import json
import re
from pathlib import Path
from typing import Any

from PIL import Image

WORDS_PER_ROW = 3
PART_GAP_MM = 25
# 抬高版本 → 已生成的板书 Word 会因指纹变化全部重建（读音改顶横线）
BOARD_DOCX_FORMAT_VERSION = "pitch-overline-v4"
# OJAD / 词典未命中时的板书读音格文案
BOARD_PITCH_NOT_FOUND_LABEL = "暂时没有在词典里面查到该词"

# Mac 常见日文字体
_FONT_CANDIDATES = (
    "/System/Library/Fonts/ヒラギノ角ゴシック W4.ttc",
    "/System/Library/Fonts/Hiragino Sans GB.ttc",
    "/System/Library/Fonts/PingFang.ttc",
    "/Library/Fonts/Arial Unicode.ttf",
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
)

_kakasi_converter = None


def to_hiragana(text: str) -> str:
    """汉字/片假名 → 平假名（板书只展示假名）。"""
    raw = (text or "").strip()
    if not raw:
        return ""
    global _kakasi_converter
    try:
        import pykakasi

        if _kakasi_converter is None:
            _kakasi_converter = pykakasi.kakasi()
        parts = _kakasi_converter.convert(raw)
        out = "".join((p.get("hira") or p.get("orig") or "") for p in parts)
        return out.strip() or raw
    except Exception:
        # 仅片假名时简单降写
        chars: list[str] = []
        for ch in raw:
            o = ord(ch)
            if 0x30A1 <= o <= 0x30F6:
                chars.append(chr(o - 0x60))
            else:
                chars.append(ch)
        return "".join(chars)


def parse_lesson_content(raw: str) -> list[str]:
    text = (raw or "").strip()
    if not text:
        return []
    # 与 jp-lesson-shared parseLessonContent 对齐：逗号/顿号/中文逗号
    parts = re.split(r"[,，、\n]+", text)
    return [p.strip() for p in parts if p.strip()]


def parse_pitch_json(raw: str | None) -> dict[str, Any] | None:
    text = (raw or "").strip()
    if not text:
        return None
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        return None
    if not isinstance(data, dict):
        return None
    kana = str(data.get("kana") or "").strip()
    pattern = str(data.get("pattern") or "").strip()
    moras_raw = data.get("moras")
    if not kana or not isinstance(moras_raw, list) or not moras_raw:
        return None
    moras: list[dict[str, str]] = []
    for item in moras_raw:
        if not isinstance(item, dict):
            continue
        c = str(item.get("c") or "").strip()
        p = str(item.get("p") or "").strip().upper()
        if not c or p not in {"L", "H", "N"}:
            continue
        moras.append({"c": c, "p": p})
    if not moras:
        return None
    return {
        "kana": kana,
        "pattern": pattern or "".join(m["p"] for m in moras),
        "moras": moras,
    }


def pitch_digest(pitch_accent: str | None, pitch_source: str | None) -> str:
    src = (pitch_source or "").strip()
    if src == "OJAD_NONE":
        return "OJAD_NONE"
    parsed = parse_pitch_json(pitch_accent)
    if not parsed:
        return ""
    return f"{parsed['kana']}|{parsed['pattern']}"


def _load_jp_font(size: int):
    from PIL import ImageFont

    for path in _FONT_CANDIDATES:
        p = Path(path)
        if not p.is_file():
            continue
        try:
            return ImageFont.truetype(str(p), size=size, index=0)
        except OSError:
            try:
                return ImageFont.truetype(str(p), size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def render_ojad_pitch_reading_png(
    pitch_accent: str | None,
    *,
    reading: str | None = None,
    word: str | None = None,
    font_size: int = 32,
) -> bytes | None:
    """画 OJAD 式顶横线读音图：连续高音一段横线；N 核音红线；只显示假名。"""
    from PIL import ImageDraw

    display = to_hiragana((reading or "").strip() or (word or "").strip())
    parsed = parse_pitch_json(pitch_accent)
    moras: list[dict[str, str]] = []
    if parsed and parsed.get("moras"):
        # 音调拍须与展示假名长度一致，否则只用假名不画错线
        joined = "".join(m["c"] for m in parsed["moras"])
        target = display or to_hiragana(parsed.get("kana") or "")
        if target and _kana_eq(joined, target):
            # 用展示假名字形，保留 OJAD 的 L/H/N
            cursor = 0
            for m in parsed["moras"]:
                n = len(m["c"])
                moras.append({"c": target[cursor : cursor + n], "p": m["p"]})
                cursor += n
            display = target
        elif not display:
            moras = list(parsed["moras"])
            display = "".join(m["c"] for m in moras)

    if not moras:
        # 无可用音调：不画裸假名，由 Word 单元格显示「暂时没有在词典里面查到该词」
        return None

    font = _load_jp_font(font_size)
    probe = Image.new("RGB", (8, 8), "white")
    pdraw = ImageDraw.Draw(probe)
    mora_widths: list[int] = []
    mora_heights: list[int] = []
    for m in moras:
        bbox = pdraw.textbbox((0, 0), m["c"], font=font)
        mora_widths.append(max(int(font_size * 1.05), bbox[2] - bbox[0] + 4))
        mora_heights.append(bbox[3] - bbox[1])
    line_gap = max(4, font_size // 6)
    bar_h = max(3, font_size // 10)
    text_h = max(mora_heights) if mora_heights else font_size
    total_w = sum(mora_widths) + 12
    total_h = line_gap + bar_h + 2 + text_h + 10
    img = Image.new("RGB", (total_w, total_h), "white")
    draw = ImageDraw.Draw(img)
    x = 6
    text_y = line_gap + bar_h + 2
    xs: list[int] = []
    for mw in mora_widths:
        xs.append(x)
        x += mw

    # 字
    for m, mw, x0 in zip(moras, mora_widths, xs):
        bbox = draw.textbbox((0, 0), m["c"], font=font)
        cw = bbox[2] - bbox[0]
        cx = x0 + (mw - cw) // 2
        color = "#e85d6f" if m["p"] == "N" else "#222222"
        draw.text((cx, text_y - bbox[1]), m["c"], font=font, fill=color)

    # 连续高音（H/N）画成一条横线；核音段用红色
    y0 = line_gap
    i = 0
    while i < len(moras):
        if moras[i]["p"] not in {"H", "N"}:
            i += 1
            continue
        j = i
        while j < len(moras) and moras[j]["p"] in {"H", "N"}:
            j += 1
        # 整段底色黑线
        x_left = xs[i]
        x_right = xs[j - 1] + mora_widths[j - 1]
        draw.rectangle([x_left, y0, x_right, y0 + bar_h], fill="#222222")
        # 核音 mora 覆盖红线
        for k in range(i, j):
            if moras[k]["p"] == "N":
                draw.rectangle(
                    [xs[k], y0, xs[k] + mora_widths[k], y0 + bar_h],
                    fill="#e85d6f",
                )
        i = j
    return _png_bytes(img)


def _kana_eq(a: str, b: str) -> bool:
    def norm(s: str) -> str:
        out: list[str] = []
        for ch in (s or "").replace(" ", ""):
            o = ord(ch)
            if 0x30A1 <= o <= 0x30F6:
                out.append(chr(o - 0x60))
            else:
                out.append(ch)
        return "".join(out)

    return norm(a) == norm(b)


def reading_label_for_cell(
    word: str,
    pitch_accent: str | None,
    reading: str | None = None,
    *,
    pitch_accent_source: str | None = None,
) -> str:
    """无图画时的纯文字兜底。

    有合法音调 → 假名；词典未查到（OJAD_NONE / 无 pitch）→ 提示文案。
    绝不写 NLLL／头高。
    """
    src = (pitch_accent_source or "").strip()
    if src == "OJAD_NONE":
        return BOARD_PITCH_NOT_FOUND_LABEL
    parsed = parse_pitch_json(pitch_accent)
    if parsed and parsed.get("kana"):
        display = to_hiragana((reading or "").strip() or (word or "").strip())
        joined = "".join(m["c"] for m in parsed["moras"])
        target = display or to_hiragana(str(parsed["kana"]))
        if target and _kana_eq(joined, target):
            return target
        return to_hiragana(str(parsed["kana"]))
    return BOARD_PITCH_NOT_FOUND_LABEL


def fnv1a_hex(text: str) -> str:
    """与 src/lib/jp-lesson-board-docx.ts buildJpLessonBoardDocxFingerprint 一致。"""
    h = 0x811C9DC5
    for ch in text:
        h ^= ord(ch)
        h = (h * 0x01000193) & 0xFFFFFFFF
    return f"v2-{h:08x}"


def build_fingerprint(
    *,
    ref_updated_at: str,
    content: str,
    meanings: str | None,
    pitch_digests: list[str],
) -> str:
    meanings_s = (meanings or "").strip()
    pitches = "\n".join((d or "").strip() for d in pitch_digests)
    raw = "\n---\n".join(
        [
            BOARD_DOCX_FORMAT_VERSION,
            (ref_updated_at or "").strip(),
            (content or "").strip(),
            meanings_s,
            pitches,
        ]
    )
    return fnv1a_hex(raw)


def _png_bytes(img: Image.Image) -> bytes:
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _row_ink_fraction(px, w: int, y: int) -> float:
    x0 = max(0, int(w * 0.04))
    x1 = min(w, int(w * 0.96))
    ink = total = 0
    for x in range(x0, x1, 2):
        r, g, b = px[x, y][:3]
        total += 1
        if not (r > 245 and g > 245 and b > 245):
            ink += 1
    return ink / total if total else 0.0


def _mean_ink(px, w: int, h: int, y0: int, y1: int) -> float:
    a = max(0, y0)
    b = min(h, y1)
    if b <= a:
        return 0.0
    return sum(_row_ink_fraction(px, w, y) for y in range(a, b)) / (b - a)


def _dark_frac(px, w: int, h: int, y0: int, y1: int) -> float:
    a = max(0, y0)
    b = min(h, y1)
    if b <= a:
        return 0.0
    x0 = max(0, int(w * 0.04))
    x1 = min(w, int(w * 0.96))
    dark = total = 0
    for y in range(a, b):
        for x in range(x0, x1, 4):
            r, g, b_ = px[x, y][:3]
            total += 1
            if (r + g + b_) / 3 < 90:
                dark += 1
    return dark / total if total else 0.0


def detect_word_card_grid_row_splits(img: Image.Image) -> list[int] | None:
    img = img.convert("RGB")
    w, h = img.size
    px = img.load()
    if w < 200 or h < 280:
        return None

    ink = [_row_ink_fraction(px, w, y) for y in range(h)]
    smooth = []
    for y in range(h):
        a = max(0, y - 1)
        b = min(h, y + 2)
        smooth.append(sum(ink[a:b]) / (b - a))

    gutters: list[tuple[int, int]] = []
    run_start = None
    for y in range(h):
        on = smooth[y] < 0.08
        if on:
            if run_start is None:
                run_start = y
        elif run_start is not None:
            if y - run_start >= 3:
                gutters.append((run_start, y - 1))
            run_start = None
    if run_start is not None and h - run_start >= 3:
        gutters.append((run_start, h - 1))

    splits: list[int] = []
    for start, end in gutters:
        mid = (start + end) // 2
        if mid <= 40:
            continue
        above = _mean_ink(px, w, h, mid - 90, mid - 8)
        below = _mean_ink(px, w, h, mid + 8, mid + 90)
        if above < 0.25 or below < 0.25:
            continue
        dark_above = _dark_frac(px, w, h, mid - 120, mid - 8)
        if mid < h * 0.28 and dark_above > 0.25:
            continue
        if splits and mid - splits[-1] < 80:
            continue
        splits.append(mid)

    if len(splits) < 1:
        return None

    if len(splits) >= 2:
        edges = [0, *splits, h]
        heights = []
        for i in range(1, len(edges)):
            if i == 1:
                continue
            heights.append(edges[i] - edges[i - 1])
        if len(heights) >= 2:
            avg = sum(heights) / len(heights)
            weird = [hh for hh in heights if hh < avg * 0.45 or hh > avg * 1.85]
            if len(weird) > len(heights) / 2:
                return None
    return splits


def card_grid_splits_to_bounds(splits: list[int], height: int) -> list[tuple[int, int]]:
    if not splits:
        return [(0, height)]
    bounds: list[tuple[int, int]] = []
    y0 = 0
    for y in splits:
        y1 = max(y0 + 1, min(height, y))
        bounds.append((y0, y1))
        y0 = y1
    if y0 < height:
        bounds.append((y0, height))
    return bounds


def group_sections_into_pages(sections: list[tuple[int, int]]) -> list[list[tuple[int, int]]]:
    """两行一块页（与 exportJpVocabRefPaginatedDocx 一致）。"""
    pages: list[list[tuple[int, int]]] = []
    i = 0
    while i < len(sections):
        pages.append(sections[i : i + 2])
        i += 2
    return pages or [[(0, 1)]]


def build_board_docx_bytes(
    *,
    image: Image.Image,
    words: list[dict[str, Any]],
) -> bytes:
    """words: [{word, pitch_accent?, reading?}, ...] 按课 content 顺序。"""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Mm, Pt
    except ImportError as exc:
        raise SystemExit(
            "需要 python-docx。请先: bash scripts/setup-jp-lesson-board-docx-mac.sh"
        ) from exc

    img = image.convert("RGB")
    w, h = img.size
    splits = detect_word_card_grid_row_splits(img)
    if splits is None:
        # 整图一页兜底
        sections = [(0, h)]
    else:
        sections = card_grid_splits_to_bounds(splits, h)

    pages = group_sections_into_pages(sections)
    surfaces = [str(x.get("word") or "").strip() for x in words]
    # 按行对齐：第 i 行对应 words[i*3 : i*3+3]
    row_words: list[list[dict[str, Any]]] = []
    for ri in range(len(sections)):
        start = ri * WORDS_PER_ROW
        row_words.append(words[start : start + WORDS_PER_ROW])

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(12)
    section.bottom_margin = Mm(12)
    section.left_margin = Mm(12)
    section.right_margin = Mm(12)

    usable_w_mm = 210 - 24  # A4 approx minus margins
    max_img_w = Mm(usable_w_mm)

    for page_idx, page_secs in enumerate(pages):
        if page_idx > 0:
            doc.add_page_break()
        for part_idx, (y0, y1) in enumerate(page_secs):
            if part_idx > 0:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Mm(PART_GAP_MM)
                p.paragraph_format.space_after = Mm(4)

            crop = img.crop((0, y0, w, y1))
            # 缩放：宽贴页，高按比例
            crop_w, crop_h = crop.size
            target_w_mm = usable_w_mm
            target_h_mm = target_w_mm * (crop_h / max(crop_w, 1))
            # 两行页时限制高度
            max_h_mm = 110 if len(page_secs) > 1 else 180
            if target_h_mm > max_h_mm:
                scale = max_h_mm / target_h_mm
                target_w_mm *= scale
                target_h_mm = max_h_mm

            png = _png_bytes(crop)
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_img.add_run()
            run.add_picture(io.BytesIO(png), width=Mm(target_w_mm))

            # 图下读音表：与词卡列对齐
            sec_index = sections.index((y0, y1)) if (y0, y1) in sections else 0
            # 更稳：按 page 内全局 section 序号
            global_sec = 0
            for pi in range(page_idx):
                global_sec += len(pages[pi])
            global_sec += part_idx
            cells_src = row_words[global_sec] if global_sec < len(row_words) else []
            cols = max(len(cells_src), 1)
            table = doc.add_table(rows=1, cols=cols)
            table.autofit = True
            for ci in range(cols):
                cell = table.rows[0].cells[ci]
                # 清空默认段落
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if ci >= len(cells_src):
                    continue
                item = cells_src[ci]
                word = str(item.get("word") or "")
                pitch = item.get("pitch_accent")
                reading = item.get("reading")
                pitch_src = item.get("pitch_accent_source")
                png = render_ojad_pitch_reading_png(
                    pitch if isinstance(pitch, str) else None,
                    reading=reading if isinstance(reading, str) else None,
                    word=word,
                )
                if png:
                    run = p.add_run()
                    # 读音图宽度随列数缩放
                    run.add_picture(io.BytesIO(png), width=Mm(min(48, usable_w_mm / cols - 4)))
                else:
                    label = reading_label_for_cell(
                        word,
                        pitch if isinstance(pitch, str) else None,
                        reading if isinstance(reading, str) else None,
                        pitch_accent_source=(
                            pitch_src if isinstance(pitch_src, str) else None
                        ),
                    )
                    run = p.add_run(label)
                    run.font.size = Pt(11 if label == BOARD_PITCH_NOT_FOUND_LABEL else 14)
                    run.font.name = "PingFang SC"
                    r = run._element
                    r.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def dry_run_lesson_148_fixture(fixture_path: Path, out_path: Path) -> dict[str, Any]:
    """本地 dry-run：用 fixture 图 + 假词表生成 docx。"""
    img = Image.open(fixture_path)
    # lesson-148 样例约 18 词 / 6 行（与切点一致）
    demo_words = [
        {"word": f"词{i+1}", "reading": "あ", "pitch_accent": None}
        for i in range(18)
    ]
    # 给前几个假 OJAD（いっぱい 头高 NLLL）
    demo_words[0] = {
        "word": "いっぱい",
        "reading": "いっぱい",
        "pitch_accent": json.dumps(
            {
                "kana": "いっぱい",
                "pattern": "NLLL",
                "moras": [
                    {"c": "い", "p": "N"},
                    {"c": "っ", "p": "L"},
                    {"c": "ぱ", "p": "L"},
                    {"c": "い", "p": "L"},
                ],
            },
            ensure_ascii=False,
        ),
    }
    demo_words[1] = {
        "word": "すごい",
        "reading": "すごい",
        "pitch_accent": json.dumps(
            {
                "kana": "すごい",
                "pattern": "LNL",
                "moras": [
                    {"c": "す", "p": "L"},
                    {"c": "ご", "p": "N"},
                    {"c": "い", "p": "L"},
                ],
            },
            ensure_ascii=False,
        ),
    }
    demo_words[6] = {
        "word": "お元気で",
        "reading": None,
        "pitch_accent": None,
        "pitch_accent_source": "OJAD_NONE",
    }
    demo_words[7] = {
        "word": "お気をつけて",
        "reading": None,
        "pitch_accent": None,
        "pitch_accent_source": "OJAD_NONE",
    }
    # 错配 pitch 不得画成「おきにいり」：假名不一致时只显示课表假名
    demo_words[8] = {
        "word": "さようなら",
        "reading": None,
        "pitch_accent": json.dumps(
            {
                "kana": "おきにいり",
                "pattern": "LHHHH",
                "moras": [
                    {"c": "お", "p": "L"},
                    {"c": "き", "p": "H"},
                    {"c": "に", "p": "H"},
                    {"c": "い", "p": "H"},
                    {"c": "り", "p": "H"},
                ],
            },
            ensure_ascii=False,
        ),
    }
    splits = detect_word_card_grid_row_splits(img)
    blob = build_board_docx_bytes(image=img, words=demo_words)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_bytes(blob)
    # 顺带写出读音预览 PNG，便于目视回归
    preview = render_ojad_pitch_reading_png(
        demo_words[0]["pitch_accent"], reading="いっぱい", word="いっぱい"
    )
    if preview:
        (out_path.parent / "lesson-148-pitch-preview.png").write_bytes(preview)
    preview2 = render_ojad_pitch_reading_png(
        demo_words[1]["pitch_accent"], reading="すごい", word="すごい"
    )
    if preview2:
        (out_path.parent / "lesson-148-pitch-preview-sugoi.png").write_bytes(preview2)
    preview_kanji = render_ojad_pitch_reading_png(
        None, reading=None, word="お気をつけて"
    )
    # 无音调应返回 None；文案走 reading_label_for_cell
    if preview_kanji is not None:
        raise AssertionError("expected no PNG when pitch missing")
    not_found = reading_label_for_cell(
        "お気をつけて", None, None, pitch_accent_source="OJAD_NONE"
    )
    (out_path.parent / "lesson-148-pitch-not-found.txt").write_text(
        not_found, encoding="utf-8"
    )
    fp = build_fingerprint(
        ref_updated_at="fixture",
        content=",".join(str(w["word"]) for w in demo_words),
        meanings=None,
        pitch_digests=[pitch_digest(w.get("pitch_accent"), None) for w in demo_words],
    )
    return {
        "ok": True,
        "splits": splits,
        "bytes": len(blob),
        "out": str(out_path),
        "fingerprint": fp,
        "sha256": hashlib.sha256(blob).hexdigest()[:16],
        "format": BOARD_DOCX_FORMAT_VERSION,
    }


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    fixture = root / "scripts" / "fixtures" / "jp-lesson-148-word-card-grid.png"
    out = root / "tmp" / "lesson-148-board-dryrun.docx"
    if not fixture.is_file():
        raise SystemExit(f"missing fixture: {fixture}")
    info = dry_run_lesson_148_fixture(fixture, out)
    print(json.dumps(info, ensure_ascii=False, indent=2))
